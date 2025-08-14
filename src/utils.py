import subprocess, os, re
from dotenv import load_dotenv
from time import sleep
from docx import Document
from datetime import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as WDA
from docx.enum.table import WD_ALIGN_VERTICAL as WVA
from docx.shared import Pt
from google.oauth2 import service_account
from googleapiclient.discovery import build
load_dotenv()

MESES = {
    1: "Janeiro", 2: "Fevereiro", 3: "Março", 4: "Abril",
    5: "Maio", 6: "Junho", 7: "Julho", 8: "Agosto",
    9: "Setembro", 10: "Outubro", 11: "Novembro", 12: "Dezembro"
}

Departamentos = {
    1: "TI",
    2: "People & Culture",
    3: "Desenvolvimento",
    4: "Marketing",
    5: "Comercial",
    6: "ADM/Financeiro",
    7: "Central de REL.",
    8: "Implantação",
    9: "Costumer Success",
    10: "Juridico"
}

modelos = {
    "Entrega Pacto": "entregapacto.docx",
    "Entrega Moovz": "entregamoovz.docx",
    "Entrega Personalfit": "entregapersonalfit.docx",
    "Entrega WAGI": "entregawagi.docx",
    "Entrega Fypay": "entregafypay.docx"
}

def pergunta_doc():
    while True:
        while True:
            print("Modelos de documento disponíveis:")
            for i, modelo in enumerate(modelos.keys(), start=1):
                print(f"{i} - {modelo}")
            try:
                opt = int(input(f"\nEscolha o modelo desejado [1 - {len(modelos)}]: "))
            except Exception as e:
                print(f"Erro: {e}. Digite um número válido.")
                sleep(0.5)
                continue
            break
        if 1 <= opt <= len(modelos):
            break
        else:
            print("Opção inválida. Tente novamente.")
            sleep(0.5)

    selected_model = list(modelos.keys())[opt - 1]
    return modelos[selected_model]

def pergunta_departamento():
    while True:            
        while True:
            print("\nDepartamnentos:\n")
            for key, valor in Departamentos.items():
                print(f"{key} - {valor}")
            try:
                dp = int(input(f"\nSelecione o departamento [1 - {len(Departamentos)}]: ")) 
            except Exception as e:
                print(f"Erro: {e}")
                continue
            break
        if len(Departamentos) >= dp > 0:
            print(f"Selecionado: {Departamentos[dp]}")                
            break               
        else:
            print("Valor inválido, tente novamente.")
            sleep(0.5)
    departamento = Departamentos[dp]
    return departamento

class DocumentEditor:
    def __init__(self, documento_entrada, documento_saida="termoentrega_modificado.docx"):
        self.documento_entrada = documento_entrada
        self.documento_saida = documento_saida        
        print("Carregando o modelo...")
        self.documento = Document(documento_entrada)

    def data_formatada(self):
        hoje = datetime.now()
        return f"Goiânia, {hoje.day} de {MESES[hoje.month]} de {hoje.year}"
            
    def pedir_dados_usuario(self):
        print("\nInforme os dados do colaborador:")
        def pedenumero():
            while True:
                n = input("Número de contato: ")
                n = re.sub(r'\D', '', n)
                if len(n) != 11:
                    print('Número inválido. O contato deve conter os 11 dígitos.')
                    sleep(0.5)
                else:
                    break
            return n
        return {
            "nome": input("Nome: ").upper(),
            "funcao": input("Função: ").upper(),
            "numero": pedenumero(),
            "departamento": pergunta_departamento().upper()
        }

    def formata_numero(self, dados):
        # Remove tudo que não for dígito do número de contato
        numero_limpo = re.sub(r'\D', '', dados["numero"])

        #Verifica se tem 11 digitos
        if len(numero_limpo) == 11:
            return re.sub(r'(\d{2})(\d{5})(\d{4})', r'(\1) \2-\3', numero_limpo)
        elif len(numero_limpo) == 10:
            return re.sub(r'(\d{2})(\d{4})(\d{4})', r'(\1) \2-\3', numero_limpo)
        else:
            return None
        
    def substituir_texto(self, substituicoes, incluir_data=True):
        if incluir_data:
            substituicoes["data"] = self.data_formatada()

        for paragrafo in self.documento.paragraphs:
            for run in paragrafo.runs:
                for marcador, valor in substituicoes.items():
                    if marcador in run.text:
                        run.text = run.text.replace(marcador, valor)

        for tabela in self.documento.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    for paragrafo in celula.paragraphs:
                        for run in paragrafo.runs:
                            for marcador, valor in substituicoes.items():
                                if marcador in run.text:
                                    run.text = run.text.replace(marcador, valor)

    def adicionar_linha_tabela(self, dados, indice_tabela=0):
        if not self.documento.tables:
            raise ValueError("Documento não contém tabelas")
        tabela = self.documento.tables[indice_tabela]
        nova_linha = tabela.add_row()
        if len(dados) != 3:
            raise ValueError("Dados devem conter [quantidade, descrição, ID]")
        for col_idx, valor in enumerate(dados):
            if col_idx < len(nova_linha.cells):
                celula = nova_linha.cells[col_idx]
                celula.text = str(valor)
                celula.vertical_alignment = WVA.CENTER
                for paragrafo in celula.paragraphs:
                    paragrafo.alignment = WDA.CENTER
                    for run in paragrafo.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(9)

    def adicionar_linha_mesclada(self, texto, indice_tabela=0):
        if not self.documento.tables:
            raise ValueError("Documento não contém tabelas")
        tabela = self.documento.tables[indice_tabela]
        nova_linha = tabela.add_row()
        celula_mesclada = nova_linha.cells[0]
        for celula in nova_linha.cells[1:]:
            celula_mesclada.merge(celula)
        celula_mesclada.text = texto
        for paragrafo in celula_mesclada.paragraphs:
            paragrafo.alignment = WDA.CENTER
            for run in paragrafo.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(7)
        celula_mesclada.vertical_alignment = WVA.CENTER

    def converter_para_pdf_libreoffice(self, nomedoc, pasta_saida="entrega_pdf"): 
        # Detecta o sistema operacional e define o caminho do LibreOffice
        import platform
        sistema_os = platform.system()
        print(f"🔍 Sistema operacional detectado: {sistema_os}")
        
        if sistema_os == "Windows":
            caminho_soffice = "C:\\Program Files\\LibreOffice\\program\\soffice.exe"
            print(f"🪟 Caminho Windows: {caminho_soffice}")
        else:
            # Linux/Docker - caminho padrão do LibreOffice
            caminho_soffice = "/usr/bin/libreoffice"
            print(f"🐧 Caminho Linux/Docker: {caminho_soffice}")
        
        if not os.path.exists(pasta_saida):
            os.makedirs(pasta_saida)
        
        try:
            # Verifica se o LibreOffice está disponível
            if not os.path.exists(caminho_soffice):
                raise FileNotFoundError(f"LibreOffice não encontrado em: {caminho_soffice}")
            
            print(f"🚀 Executando: {caminho_soffice}")
            subprocess.run([
                caminho_soffice,
                "--headless",
                "--norestore",
                "--nolockcheck",
                "--nofirststartwizard",
                "--nologo",
                "--convert-to", "pdf",
                "--outdir", pasta_saida,
                nomedoc
            ], check=True, capture_output=True, text=True)
            print("\n✅ Conversão concluída.")
        except subprocess.CalledProcessError as e:
            print(f"❌ Erro na conversão: {e}")
            print(f"Erro de saída: {e.stderr}")
        except FileNotFoundError as e:
            print(f"❌ {e}")
            print("💡 Certifique-se de que o LibreOffice está instalado")
        except Exception as e:
            print(f"❌ Erro inesperado na conversão: {e}")
            print(f"Tipo de erro: {type(e)}")
            print(f"Detalhes: {str(e)}")

    def salvar(self, nome_colaborador=None):
        if not os.path.exists("entrega_docx"):
            os.makedirs("entrega_docx")
        documento_saida = f"entrega_docx/Termo de entrega {nome_colaborador}.docx"
        self.documento.save(documento_saida)
        if nome_colaborador:
            self.converter_para_pdf_libreoffice(nomedoc=documento_saida) # Converte o documento para PDF usando LibreOffice
            print(f"\nDocumento salvo como: Termo de entrega {nome_colaborador}.pdf")
        else:
            print(f"\nDocumento salvo como: Termo de entrega {nome_colaborador}.docx")

class SheetsHandler:
    def __init__(self, spreadsheet_id, nome_proprietario=None):
        SERVICE_ACCOUNT_FILE = os.getenv("CREDENTIALS")  # Caminho para o arquivo JSON da conta de serviço
        SCOPES = ['https://www.googleapis.com/auth/spreadsheets']

        credentials = service_account.Credentials.from_service_account_file(
            SERVICE_ACCOUNT_FILE, scopes=SCOPES
        )

        self.service = build('sheets', 'v4', credentials=credentials)
        self.spreadsheet_id = spreadsheet_id
        self.nome_proprietario = nome_proprietario
        self.departamento = None
        
    def obter_metadados(self):
        try:
            # Usa o cliente autenticado para obter os metadados da planilha
            sheet_metadata = self.service.spreadsheets().get(spreadsheetId=self.spreadsheet_id).execute()
            return [sheet["properties"]["title"] for sheet in sheet_metadata["sheets"]]
        except Exception as e:
            print("Erro ao acessar metadados da planilha:", e)
            return None

    def buscar_palavra_em_abas(self, palavra_busca, sheet_names):
        resultados = []

        for sheet_name in sheet_names:
            try:
                result = self.service.spreadsheets().values().get(
                    spreadsheetId=self.spreadsheet_id,
                    range=sheet_name
                ).execute()

                data = result.get("values", [])
                for row_idx, row in enumerate(data):
                    for col_idx, cell in enumerate(row):
                        if palavra_busca.lower() in str(cell).lower():
                            linha_completa = data[row_idx]
                            resultados.append((sheet_name, row_idx + 1, linha_completa))
                            break  # para evitar múltiplas capturas na mesma linha

            except Exception as e:
                print(f"Erro ao acessar dados da aba '{sheet_name}': {e}")

        return resultados

    def exibir_resultados(self, resultados):
        if resultados:
            print("\n🔍 Resultados encontrados:\n")
            for aba, linha, valores in resultados:
                valores = [item for item in valores if item and item.strip()]
                self.nome_proprietario = valores[0] if len(valores) > 0 else None
                print(f"Proprietário: {self.nome_proprietario}") if self.nome_proprietario else print("Proprietário não encontrado.")
                lista_filtrada = self.filtrar_lista_por_aba(aba, valores)
                print(f"Item: {' | '.join(lista_filtrada)}\n")
        else:
            print("\n❌ ID não encontrado na planilha.\n")

    def altera_valor_planilha(self, spreadsheet_id, aba, coluna, linha, valor):
        # nome_proprietario = novo_proprietario
        SERVICE_ACCOUNT_FILE = os.getenv("CREDENTIALS")  # Caminho para o arquivo JSON da conta de serviço

        SCOPES = ['https://www.googleapis.com/auth/spreadsheets']

        credentials = service_account.Credentials.from_service_account_file(
            SERVICE_ACCOUNT_FILE, scopes=SCOPES
        )

        service = build('sheets', 'v4', credentials=credentials)
        
        # Define a célula a ser atualizada usando notação A1
        range_update = f"{aba}!{coluna}{linha}"

        # Corpo da requisição: lista de listas
        body = {
            "values": [[valor]]
        }

        try:
            result = service.spreadsheets().values().update(
                spreadsheetId=spreadsheet_id,
                range=range_update,
                valueInputOption="RAW",
                body=body
            ).execute()
            
            print("✔️ Proprietário atualizado na planilha.") if coluna == "a" else print("✔️ Departamento atualizado na planilha.")
            
        except Exception as e:
            print(f"\n❌ Erro ao atualizar valor: {e}")

    def altera_proprietario(self, spreadsheet_id, aba, linha, novo_proprietario):
        self.altera_valor_planilha(spreadsheet_id, aba, "a", linha, novo_proprietario)
            
    def altera_departamento(self, spreadsheet_id, aba, linha, novo_departamento):
        self.altera_valor_planilha(spreadsheet_id, aba, "b", linha, novo_departamento)
    
    @staticmethod
    def filtrar_lista_por_aba(aba, valores):
        # Remove todos os valores vazios/nulos da linha original
        valores_limpos = [str(v).strip() for v in valores if v and str(v).strip()]
        
        # Define os índices a serem filtrados para cada aba
        indices_por_aba = {
            "HEADSET": [2, 3, 4, 5],  # ['MARCA', 'MODELO', 'CONEXÃO', 'ID']
            "DESKTOP'S": [3, 4, 5, 6, 7, 8, 9], 
            "CELULARES-TABLETS": [3, 4, 5, 6],
            "MONITORES": [3, 4, 5, 6],
            "IMP-TRITURADORA": [3, 4, 5, 6],
            "NOTEBOOKS": [3, 4, 5, 6, 7, 8, 9],
            "OUTROS": [1, 4, 5, 6]
        }
        
        indices = indices_por_aba.get(aba, [])
        return [valores_limpos[i] for i in indices if i < len(valores_limpos)]